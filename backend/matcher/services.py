"""
AI and external service integrations for the matcher app.
"""

import os
import json
import logging
import time
from typing import Dict, Any, Optional, List
from pathlib import Path

import google.generativeai as genai
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile

import PyPDF2
import docx
from io import BytesIO

from .exceptions import GeminiAPIException, FileProcessingException, ConfigurationException
from .logging_config import ai_logger, performance_logger
from .error_recovery import (
    gemini_retry_config, gemini_circuit_breaker, gemini_bulkhead,
    retry_with_backoff, safe_execute, with_performance_monitoring
)

logger = logging.getLogger(__name__)

# For backward compatibility
GeminiAPIError = GeminiAPIException


class GeminiResumeParser:
    """
    Google Gemini API client for resume parsing and analysis with comprehensive error handling
    """
    
    def __init__(self):
        self.api_key = getattr(settings, 'GEMINI_API_KEY', '')
        self.model_name = getattr(settings, 'GEMINI_MODEL_NAME', 'gemini-pro')
        
        if not self.api_key:
            raise ConfigurationException(
                message="GEMINI_API_KEY not configured in settings",
                code="GEMINI_API_KEY_MISSING"
            )
        
        try:
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel(self.model_name)
            logger.info(f"Gemini model {self.model_name} initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize Gemini model: {str(e)}")
            raise GeminiAPIException(
                message=f"Failed to initialize Gemini model: {str(e)}",
                code="GEMINI_INITIALIZATION_ERROR"
            )
    
    @gemini_circuit_breaker
    @gemini_bulkhead
    @retry_with_backoff(gemini_retry_config)
    @with_performance_monitoring
    def parse_resume(self, file_path: str, file_content: Optional[bytes] = None) -> Dict[str, Any]:
        """
        Parse resume using Google Gemini API with comprehensive error handling
        
        Args:
            file_path: Path to the resume file
            file_content: Optional file content as bytes
            
        Returns:
            Dict containing parsed resume data
        """
        start_time = time.time()
        file_size = len(file_content) if file_content else 0
        
        try:
            # Extract text from file
            text_content = self._extract_text_from_file(file_path, file_content)
            
            if not text_content.strip():
                raise FileProcessingException(
                    message="No text content could be extracted from the file",
                    code="EMPTY_FILE_CONTENT"
                )
            
            # Generate structured data using Gemini
            parsed_data = self._generate_structured_data(text_content)
            
            processing_time = time.time() - start_time
            
            # Log successful operation
            ai_logger.log_gemini_request(
                operation='resume_parse',
                input_size=file_size or len(text_content.encode('utf-8')),
                processing_time=processing_time,
                success=True,
                confidence_score=parsed_data.get('confidence_score', 0.8),
                file_path=file_path
            )
            
            return {
                'success': True,
                'parsed_text': text_content,
                'structured_data': parsed_data,
                'processing_time': processing_time,
                'confidence_score': parsed_data.get('confidence_score', 0.8),
                'timestamp': time.time()
            }
            
        except (GeminiAPIException, FileProcessingException) as e:
            processing_time = time.time() - start_time
            
            # Log failed operation
            ai_logger.log_gemini_request(
                operation='resume_parse',
                input_size=file_size,
                processing_time=processing_time,
                success=False,
                error=str(e),
                file_path=file_path
            )
            
            raise e
            
        except Exception as e:
            processing_time = time.time() - start_time
            
            # Log unexpected error
            ai_logger.log_gemini_request(
                operation='resume_parse',
                input_size=file_size,
                processing_time=processing_time,
                success=False,
                error=f"Unexpected error: {str(e)}",
                file_path=file_path
            )
            
            raise GeminiAPIException(
                message=f"Unexpected error during resume parsing: {str(e)}",
                code="GEMINI_UNEXPECTED_ERROR"
            )
    
    def _extract_text_from_file(self, file_path: str, file_content: Optional[bytes] = None) -> str:
        """
        Extract text content from PDF, DOCX, or TXT files with error handling
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_content:
                # Use provided file content
                if file_extension == '.pdf':
                    return self._extract_text_from_pdf_bytes(file_content)
                elif file_extension in ['.doc', '.docx']:
                    return self._extract_text_from_docx_bytes(file_content)
                elif file_extension == '.txt':
                    return file_content.decode('utf-8', errors='ignore')
            else:
                # Read from file path
                if file_extension == '.pdf':
                    return self._extract_text_from_pdf_file(file_path)
                elif file_extension in ['.doc', '.docx']:
                    return self._extract_text_from_docx_file(file_path)
                elif file_extension == '.txt':
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        return f.read()
            
            raise FileProcessingException(
                message=f"Unsupported file format: {file_extension}",
                code="UNSUPPORTED_FILE_FORMAT"
            )
            
        except FileProcessingException:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise FileProcessingException(
                message=f"Failed to extract text: {str(e)}",
                code="TEXT_EXTRACTION_ERROR"
            )
    
    def _extract_text_from_pdf_file(self, file_path: str) -> str:
        """Extract text from PDF file with error handling"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise FileProcessingException(
                message=f"Error reading PDF file: {str(e)}",
                code="PDF_READING_ERROR"
            )
        
        return text.strip()
    
    def _extract_text_from_pdf_bytes(self, file_content: bytes) -> str:
        """Extract text from PDF bytes with error handling"""
        text = ""
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            raise FileProcessingException(
                message=f"Error reading PDF content: {str(e)}",
                code="PDF_CONTENT_ERROR"
            )
        
        return text.strip()
    
    def _extract_text_from_docx_file(self, file_path: str) -> str:
        """Extract text from DOCX file with error handling"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            raise FileProcessingException(
                message=f"Error reading DOCX file: {str(e)}",
                code="DOCX_READING_ERROR"
            )
        
        return text.strip()
    
    def _extract_text_from_docx_bytes(self, file_content: bytes) -> str:
        """Extract text from DOCX bytes with error handling"""
        text = ""
        try:
            docx_file = BytesIO(file_content)
            doc = docx.Document(docx_file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            raise FileProcessingException(
                message=f"Error reading DOCX content: {str(e)}",
                code="DOCX_CONTENT_ERROR"
            )
        
        return text.strip()
    
    def _generate_structured_data(self, text_content: str) -> Dict[str, Any]:
        """
        Use Gemini API to generate structured data from resume text with error handling
        """
        try:
            prompt = self._create_parsing_prompt(text_content)
            
            # Generate content using Gemini
            response = self.model.generate_content(prompt)
            
            if not response or not hasattr(response, 'text'):
                raise GeminiAPIException(
                    message="Invalid response from Gemini API",
                    code="GEMINI_INVALID_RESPONSE"
                )
            
            # Parse the JSON response
            try:
                structured_data = json.loads(response.text)
            except json.JSONDecodeError:
                # If JSON parsing fails, try to extract JSON from the response
                structured_data = self._extract_json_from_response(response.text)
            
            # Validate and clean the structured data
            return self._validate_structured_data(structured_data)
            
        except GeminiAPIException:
            raise
        except Exception as e:
            logger.error(f"Error generating structured data: {str(e)}")
            raise GeminiAPIException(
                message=f"Failed to generate structured data: {str(e)}",
                code="GEMINI_STRUCTURED_DATA_ERROR"
            )
    
    def _create_parsing_prompt(self, text_content: str) -> str:
        """
        Create a detailed prompt for resume parsing
        """
        return f"""
Please analyze the following resume text and extract structured information in JSON format.

Resume Text:
{text_content}

Please extract the following information and return it as a valid JSON object:

{{
    "personal_info": {{
        "name": "Full name of the candidate",
        "email": "Email address",
        "phone": "Phone number",
        "location": "Current location/address",
        "linkedin": "LinkedIn profile URL",
        "github": "GitHub profile URL",
        "portfolio": "Portfolio website URL"
    }},
    "summary": "Professional summary or objective (2-3 sentences)",
    "experience": [
        {{
            "company": "Company name",
            "position": "Job title",
            "duration": "Employment duration",
            "description": "Job description and achievements",
            "technologies": ["List of technologies used"]
        }}
    ],
    "education": [
        {{
            "institution": "School/University name",
            "degree": "Degree type and field",
            "duration": "Study period",
            "gpa": "GPA if mentioned"
        }}
    ],
    "skills": {{
        "technical_skills": ["List of technical skills"],
        "programming_languages": ["Programming languages"],
        "frameworks": ["Frameworks and libraries"],
        "tools": ["Tools and software"],
        "soft_skills": ["Soft skills"]
    }},
    "certifications": [
        {{
            "name": "Certification name",
            "issuer": "Issuing organization",
            "date": "Date obtained"
        }}
    ],
    "projects": [
        {{
            "name": "Project name",
            "description": "Project description",
            "technologies": ["Technologies used"],
            "url": "Project URL if available"
        }}
    ],
    "languages": [
        {{
            "language": "Language name",
            "proficiency": "Proficiency level"
        }}
    ],
    "total_experience_years": "Estimated total years of experience as a number",
    "confidence_score": "Confidence score between 0.0 and 1.0 for the extraction quality"
}}

Important:
- Return only valid JSON, no additional text or formatting
- If information is not available, use null or empty arrays/objects
- Be as accurate as possible in extracting information
- For experience years, provide your best estimate as a number
- Confidence score should reflect how well you could extract the information
"""
    
    def _extract_json_from_response(self, response_text: str) -> Dict[str, Any]:
        """
        Try to extract JSON from response text if direct parsing fails
        """
        try:
            # Look for JSON content between curly braces
            import re
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                return json.loads(json_str)
            else:
                # Fallback: create basic structure
                return self._create_fallback_structure()
        except Exception:
            return self._create_fallback_structure()
    
    def _create_fallback_structure(self) -> Dict[str, Any]:
        """
        Create a fallback structure when parsing fails
        """
        return {
            "personal_info": {
                "name": None,
                "email": None,
                "phone": None,
                "location": None,
                "linkedin": None,
                "github": None,
                "portfolio": None
            },
            "summary": None,
            "experience": [],
            "education": [],
            "skills": {
                "technical_skills": [],
                "programming_languages": [],
                "frameworks": [],
                "tools": [],
                "soft_skills": []
            },
            "certifications": [],
            "projects": [],
            "languages": [],
            "total_experience_years": 0,
            "confidence_score": 0.1
        }
    
    def _validate_structured_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Validate and clean the structured data
        """
        # Ensure required keys exist
        required_keys = [
            'personal_info', 'summary', 'experience', 'education', 
            'skills', 'certifications', 'projects', 'languages',
            'total_experience_years', 'confidence_score'
        ]
        
        for key in required_keys:
            if key not in data:
                if key in ['experience', 'education', 'certifications', 'projects', 'languages']:
                    data[key] = []
                elif key == 'skills':
                    data[key] = {
                        "technical_skills": [],
                        "programming_languages": [],
                        "frameworks": [],
                        "tools": [],
                        "soft_skills": []
                    }
                elif key == 'personal_info':
                    data[key] = {}
                elif key == 'total_experience_years':
                    data[key] = 0
                elif key == 'confidence_score':
                    data[key] = 0.5
                else:
                    data[key] = None
        
        # Validate confidence score
        if not isinstance(data['confidence_score'], (int, float)) or data['confidence_score'] < 0 or data['confidence_score'] > 1:
            data['confidence_score'] = 0.5
        
        # Validate experience years
        if not isinstance(data['total_experience_years'], (int, float)) or data['total_experience_years'] < 0:
            data['total_experience_years'] = 0
        
        return data


class FileValidator:
    """
    File validation service for resume uploads
    """
    
    ALLOWED_EXTENSIONS = ['.pdf', '.doc', '.docx', '.txt']
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @classmethod
    def validate_file(cls, file) -> Dict[str, Any]:
        """
        Validate uploaded file
        
        Args:
            file: Django UploadedFile object
            
        Returns:
            Dict with validation results
        """
        errors = []
        
        # Check file extension
        file_extension = Path(file.name).suffix.lower()
        if file_extension not in cls.ALLOWED_EXTENSIONS:
            errors.append(f"File type {file_extension} not supported. Allowed types: {', '.join(cls.ALLOWED_EXTENSIONS)}")
        
        # Check file size
        if file.size > cls.MAX_FILE_SIZE:
            errors.append(f"File size ({file.size} bytes) exceeds maximum allowed size ({cls.MAX_FILE_SIZE} bytes)")
        
        # Check if file is empty
        if file.size == 0:
            errors.append("File is empty")
        
        # Basic file content validation
        try:
            # Read a small portion to check if file is readable
            file.seek(0)
            sample = file.read(1024)
            file.seek(0)
            
            if not sample:
                errors.append("File appears to be empty or corrupted")
        except Exception as e:
            errors.append(f"File reading error: {str(e)}")
        
        return {
            'is_valid': len(errors) == 0,
            'errors': errors,
            'file_extension': file_extension,
            'file_size': file.size
        }
    
    @classmethod
    def sanitize_filename(cls, filename: str) -> str:
        """
        Sanitize filename for safe storage
        """
        import re
        import uuid
        
        # Get file extension
        file_extension = Path(filename).suffix.lower()
        
        # Remove extension and clean filename
        name_without_ext = Path(filename).stem
        
        # Remove special characters and spaces
        clean_name = re.sub(r'[^\w\-_.]', '_', name_without_ext)
        
        # Limit length
        clean_name = clean_name[:50]
        
        # Add timestamp to avoid conflicts
        timestamp = str(int(time.time()))
        
        return f"{clean_name}_{timestamp}{file_extension}"